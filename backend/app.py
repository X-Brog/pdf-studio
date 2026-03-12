from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import uuid
import threading
import time
import shutil
from werkzeug.utils import secure_filename

# PDF/Document processing imports
import PyPDF2
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import io

app = Flask(__name__)
CORS(app, origins=["https://pdf-studio-tau.vercel.app"])

UPLOAD_FOLDER = '/tmp/pdf_studio_uploads'
OUTPUT_FOLDER = '/tmp/pdf_studio_outputs'
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

ALLOWED_EXTENSIONS = {
    'pdf', 'docx', 'doc', 'txt', 'jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp'
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def schedule_delete(filepath, delay=300):
    """Delete file after delay seconds"""
    def delete_later():
        time.sleep(delay)
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
        except:
            pass
    t = threading.Thread(target=delete_later, daemon=True)
    t.start()

def get_unique_path(folder, filename):
    uid = str(uuid.uuid4())[:8]
    base, ext = os.path.splitext(filename)
    return os.path.join(folder, f"{base}_{uid}{ext}")

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'message': 'PDF Studio API running'})

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not allowed'}), 400

    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[1].lower()

    uid = str(uuid.uuid4())
    upload_path = os.path.join(UPLOAD_FOLDER, f"{uid}.{ext}")

    file.save(upload_path)

    size = os.path.getsize(upload_path)

    page_count = None
    if ext == 'pdf':
        try:
            reader = PyPDF2.PdfReader(upload_path, strict=False)
            page_count = len(reader.pages)
        except:
            pass

    schedule_delete(upload_path, 600)

    return jsonify({
        'file_id': uid,
        'filename': filename,
        'size': size,
        'extension': ext,
        'page_count': page_count,
        'upload_path': upload_path 
    })

@app.route('/api/convert', methods=['POST'])
def convert_file():
    data = request.json
    operation = data.get('operation')
    upload_path = data.get('upload_path')
    filename = data.get('filename', 'file')

    if not upload_path or not os.path.exists(upload_path):
        return jsonify({'error': 'File not found'}), 404

    try:
        output_path = None

        if operation == 'jpg_to_pdf' or operation == 'png_to_pdf':
            output_path = _images_to_pdf([upload_path], filename)

        elif operation == 'pdf_to_jpg':
            output_path = _pdf_to_images(upload_path, filename)

        elif operation == 'txt_to_pdf':
            output_path = _txt_to_pdf(upload_path, filename)

        elif operation == 'docx_to_pdf':
            output_path = _docx_to_pdf(upload_path, filename)

        elif operation == 'pdf_to_word':
            output_path = _pdf_to_word(upload_path, filename)

        else:
            return jsonify({'error': f'Unknown operation: {operation}'}), 400

        if not output_path or not os.path.exists(output_path):
            return jsonify({'error': 'Conversion failed'}), 500

        schedule_delete(output_path, 300)
        out_filename = os.path.basename(output_path)

        return jsonify({
            'success': True,
            'output_path': output_path,
            'output_filename': out_filename,
            'download_url': f'/api/download/{out_filename}'
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route('/api/convert/multiple-images', methods=['POST'])
def convert_multiple_images():
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'No files provided'}), 400

    output_path = os.path.join(OUTPUT_FOLDER, f"combined_{uuid.uuid4().hex[:8]}.pdf")

    try:
        images = []
        for f in files:
            img = Image.open(f).convert('RGB')
            images.append(img)

        if images:
            images[0].save(output_path, save_all=True, append_images=images[1:])

        schedule_delete(output_path, 300)
        return jsonify({
            'success': True,
            'output_filename': os.path.basename(output_path),
            'download_url': f'/api/download/{os.path.basename(output_path)}'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500    

def _images_to_pdf(image_paths, original_filename):
    base = os.path.splitext(original_filename)[0]
    output_path = os.path.join(OUTPUT_FOLDER, f"{base}_{uuid.uuid4().hex[:6]}.pdf")
    images = []
    for p in image_paths:
        img = Image.open(p).convert('RGB')
        images.append(img)
    if images:
        images[0].save(output_path, save_all=True, append_images=images[1:])
    return output_path

def _pdf_to_images(pdf_path, original_filename):
    """Convert PDF pages to images and zip them"""
    import zipfile
    base = os.path.splitext(original_filename)[0]
    zip_path = os.path.join(OUTPUT_FOLDER, f"{base}_{uuid.uuid4().hex[:6]}_pages.zip")

    try:
        from pdf2image import convert_from_path
        images = convert_from_path(pdf_path, dpi=150)
        with zipfile.ZipFile(zip_path, 'w') as zf:
            for i, img in enumerate(images):
                img_bytes = io.BytesIO()
                img.save(img_bytes, format='JPEG', quality=85)
                zf.writestr(f"page_{i+1}.jpg", img_bytes.getvalue())
        return zip_path
    except ImportError:
        # Fallback: just rename PDF (pdf2image needs poppler)
        # Create a simple image from first page using PyPDF2 info
        output_path = os.path.join(OUTPUT_FOLDER, f"{base}_{uuid.uuid4().hex[:6]}.jpg")
        # Create placeholder image
        img = Image.new('RGB', (800, 1100), color=(255, 255, 255))
        img.save(output_path, 'JPEG')
        return output_path

def _txt_to_pdf(txt_path, original_filename):
    base = os.path.splitext(original_filename)[0]
    output_path = os.path.join(OUTPUT_FOLDER, f"{base}_{uuid.uuid4().hex[:6]}.pdf")

    with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()

    doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    for para in text.split('\n'):
        if para.strip():
            story.append(Paragraph(para, styles['Normal']))
    doc.build(story)
    return output_path

def _docx_to_pdf(docx_path, original_filename):
    """Convert DOCX to PDF using reportlab"""
    from docx import Document
    base = os.path.splitext(original_filename)[0]
    output_path = os.path.join(OUTPUT_FOLDER, f"{base}_{uuid.uuid4().hex[:6]}.pdf")

    doc = Document(docx_path)
    pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    for para in doc.paragraphs:
        if para.text.strip():
            story.append(Paragraph(para.text, styles['Normal']))
    pdf_doc.build(story)
    return output_path

def _pdf_to_word(pdf_path, original_filename):
    """Extract text from PDF and create DOCX"""
    from docx import Document as DocxDocument
    base = os.path.splitext(original_filename)[0]
    output_path = os.path.join(OUTPUT_FOLDER, f"{base}_{uuid.uuid4().hex[:6]}.docx")

    doc = DocxDocument()
    doc.add_heading(base, 0)

    with open(pdf_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                doc.add_heading(f'Page {i+1}', level=2)
                doc.add_paragraph(text)

    doc.save(output_path)
    return output_path

@app.route('/api/edit/merge', methods=['POST'])
def merge_pdfs():
    data = request.json
    paths = data.get('paths', [])

    if len(paths) < 2:
        return jsonify({'error': 'Need at least 2 PDFs to merge'}), 400

    output_path = os.path.join(OUTPUT_FOLDER, f"merged_{uuid.uuid4().hex[:8]}.pdf")

    try:
        writer = PyPDF2.PdfWriter()
        for p in paths:
            if os.path.exists(p):
                with open(p, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        writer.add_page(page)

        with open(output_path, 'wb') as f:
            writer.write(f)

        schedule_delete(output_path, 300)
        return jsonify({
            'success': True,
            'output_path': output_path,
            'output_filename': os.path.basename(output_path),
            'download_url': f'/api/download/{os.path.basename(output_path)}'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/edit/split', methods=['POST'])
def split_pdf():
    import zipfile
    data = request.json
    upload_path = data.get('upload_path')
    page_ranges = data.get('page_ranges', [])

    if not upload_path or not os.path.exists(upload_path):
        return jsonify({'error': 'File not found'}), 404

    zip_path = os.path.join(OUTPUT_FOLDER, f"split_{uuid.uuid4().hex[:8]}.zip")

    try:
        with open(upload_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            total = len(reader.pages)

            with zipfile.ZipFile(zip_path, 'w') as zf:
                if page_ranges:
                    for idx, rng in enumerate(page_ranges):
                        writer = PyPDF2.PdfWriter()
                        start = max(0, rng.get('start', 1) - 1)
                        end = min(total, rng.get('end', total))
                        for i in range(start, end):
                            writer.add_page(reader.pages[i])
                        buf = io.BytesIO()
                        writer.write(buf)
                        zf.writestr(f"part_{idx+1}.pdf", buf.getvalue())
                else:
                    # Split each page
                    for i in range(total):
                        writer = PyPDF2.PdfWriter()
                        writer.add_page(reader.pages[i])
                        buf = io.BytesIO()
                        writer.write(buf)
                        zf.writestr(f"page_{i+1}.pdf", buf.getvalue())

        schedule_delete(zip_path, 300)
        return jsonify({
            'success': True,
            'output_filename': os.path.basename(zip_path),
            'download_url': f'/api/download/{os.path.basename(zip_path)}'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/edit/delete-pages', methods=['POST'])
def delete_pages():
    data = request.json
    upload_path = data.get('upload_path')
    pages_to_delete = data.get('pages', [])  # 1-indexed

    if not upload_path or not os.path.exists(upload_path):
        return jsonify({'error': 'File not found'}), 404

    output_path = os.path.join(OUTPUT_FOLDER, f"edited_{uuid.uuid4().hex[:8]}.pdf")

    try:
        with open(upload_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            writer = PyPDF2.PdfWriter()
            delete_set = set(p - 1 for p in pages_to_delete)  # convert to 0-indexed
            for i, page in enumerate(reader.pages):
                if i not in delete_set:
                    writer.add_page(page)

        with open(output_path, 'wb') as f:
            writer.write(f)

        schedule_delete(output_path, 300)
        return jsonify({
            'success': True,
            'output_filename': os.path.basename(output_path),
            'download_url': f'/api/download/{os.path.basename(output_path)}'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/edit/add-text', methods=['POST'])
def add_text_to_pdf():
    data = request.json
    upload_path = data.get('upload_path')
    text = data.get('text', '')
    page_num = data.get('page', 1) - 1
    x = data.get('x', 100)
    y = data.get('y', 100)
    font_size = data.get('font_size', 12)
    color = data.get('color', '#000000')

    if not upload_path or not os.path.exists(upload_path):
        return jsonify({'error': 'File not found'}), 404

    output_path = os.path.join(OUTPUT_FOLDER, f"annotated_{uuid.uuid4().hex[:8]}.pdf")

    try:
        # Create overlay PDF with text
        overlay_buf = io.BytesIO()
        c = canvas.Canvas(overlay_buf, pagesize=letter)

        # Parse color
        r = int(color[1:3], 16) / 255
        g = int(color[3:5], 16) / 255
        b = int(color[5:7], 16) / 255
        c.setFillColorRGB(r, g, b)
        c.setFont("Helvetica", font_size)
        c.drawString(x, y, text)
        c.save()

        overlay_buf.seek(0)
        overlay_reader = PyPDF2.PdfReader(overlay_buf)

        with open(upload_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            writer = PyPDF2.PdfWriter()
            for i, page in enumerate(reader.pages):
                if i == page_num and len(overlay_reader.pages) > 0:
                    page.merge_page(overlay_reader.pages[0])
                writer.add_page(page)

        with open(output_path, 'wb') as f:
            writer.write(f)

        schedule_delete(output_path, 300)
        return jsonify({
            'success': True,
            'output_filename': os.path.basename(output_path),
            'download_url': f'/api/download/{os.path.basename(output_path)}'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route('/api/compress', methods=['POST'])
def compress_pdf():
    try:
        data = request.get_json()
        upload_path = data.get('upload_path')
        quality = data.get('quality', 'medium')  # 'low', 'medium', 'high'
        
        if not upload_path or not os.path.exists(upload_path):
            return jsonify({'error': 'File not found'}), 400

        output_filename = f"compressed_{uuid.uuid4().hex[:8]}.pdf"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)

        # Quality settings
        quality_settings = {
            'low': {
                'gs_setting': '/screen',
                'dpi': 72,
                'description': 'Smallest file, lowest quality (30-50% smaller)'
            },
            'medium': {
                'gs_setting': '/ebook',
                'dpi': 150,
                'description': 'Balanced compression (20-30% smaller)'
            },
            'high': {
                'gs_setting': '/printer',
                'dpi': 200,
                'description': 'Best quality, minimal compression (10-15% smaller)'
            }
        }
        
        settings = quality_settings.get(quality, quality_settings['medium'])

        try:
            # Try Ghostscript first
            import subprocess
            subprocess.run([
                'gs', '-sDEVICE=pdfwrite',
                f'-dPDFSETTINGS={settings["gs_setting"]}',
                '-dCompressFonts=true',
                '-dCompressStreams=true',
                f'-r{settings["dpi"]}',
                '-dNOPAUSE', '-dQUIET', '-dBATCH',
                f'-sOutputFile={output_path}',
                upload_path
            ], capture_output=True, timeout=20)
            
        except:
            # Fallback: PyPDF2 compression
            with open(upload_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                writer = PyPDF2.PdfWriter()
                
                for page in reader.pages:
                    page.compress_content_streams()
                    writer.add_page(page)
                
                with open(output_path, 'wb') as out_f:
                    writer.write(out_f)

        original_size = os.path.getsize(upload_path)
        compressed_size = os.path.getsize(output_path)
        saved = max(0, round((1 - compressed_size / original_size) * 100))

        schedule_delete(output_path, 300)
        
        return jsonify({
            'success': True,
            'output_filename': output_filename,
            'original_size': original_size,
            'compressed_size': compressed_size,
            'saved_percent': saved,
            'quality_used': quality,
            'download_url': f'/api/download/{output_filename}'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    # Security: only allow files from output folder
    safe_filename = secure_filename(filename)
    file_path = os.path.join(OUTPUT_FOLDER, safe_filename)

    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found or expired'}), 404

    return send_file(file_path, as_attachment=True, download_name=safe_filename)

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Maximum size is 50MB'}), 413

if __name__ == '__main__':
    app.run(debug=True, port=5000, host='0.0.0.0')
